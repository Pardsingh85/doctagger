import os
import re
import io
from docx import Document
import pdfplumber
from dotenv import load_dotenv
from openai import OpenAI  # Requires openai>=1.0.0

# Load .env variables from root (for local dev)
env_path = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '.env'))
load_dotenv(dotenv_path=env_path)

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

def extract_text(uploaded_file):
    """
    Extracts clean text from UploadFile (PDF, DOCX, or TXT).
    """
    ext = os.path.splitext(uploaded_file.filename)[1].lower()
    file_bytes = uploaded_file.file.read()

    if ext == ".txt":
        return file_bytes.decode("utf-8", errors="ignore")

    elif ext == ".docx":
        doc = Document(io.BytesIO(file_bytes))
        text_chunks = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_chunks.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_chunks.append(" | ".join(row_text))

        return "\n".join(text_chunks)

    elif ext == ".pdf":
        text = ""
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text

    return ""

def get_tags(text, mode_prompt="", num_tags=8, mode="Keywords"):
    """
    Calls OpenAI API to extract tags from text.
    Supports multiple prompt modes.
    """
    try:
        base_prompt = {
            "Keywords": f"Extract exactly {num_tags} keywords that summarize the content of this document.",
            "Topics": f"List {num_tags} major themes or subjects covered in the document.",
            "Custom Prompt": mode_prompt
        }.get(mode, "Extract relevant tags.")

        prompt = base_prompt

        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are an AI assistant that extracts tags from documents."},
                {"role": "user", "content": f"{prompt}\n\n{text}"}
            ],
            temperature=0.5
        )
        return response.choices[0].message.content.strip()

    except Exception as e:
        return f"❌ Error: {e}"

def parse_tags(raw_text):
    """
    Parses raw OpenAI output into a list of clean, deduplicated tags.
    """
    lines = raw_text.split("\n")
    tags = []

    for line in lines:
        # Remove bullets, numbers, dashes, etc.
        line = re.sub(r"^[\d\.\-\•\)]*\s*", "", line)
        parts = [t.strip() for t in line.split(",") if t.strip()]
        tags.extend(parts)

    # Deduplicate while preserving order
    seen = set()
    clean_tags = []
    for tag in tags:
        if tag.lower() not in seen and tag:
            seen.add(tag.lower())
            clean_tags.append(tag)

    return clean_tags
