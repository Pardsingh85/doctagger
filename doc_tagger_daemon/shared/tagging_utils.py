# shared/tagging_utils.py
from __future__ import annotations
import os
import re
import io
from typing import List
from .secrets import get_secret

def _openai_client():
    """
    Create OpenAI client on demand. Prefers Key Vault secret 'OpenAI-ApiKey',
    falls back to env var OPENAI_API_KEY.
    """
    key = get_secret("OpenAI-ApiKey", default=os.getenv("OPENAI_API_KEY"))
    if not key:
        raise RuntimeError("OpenAI API key not set (Key Vault 'OpenAI-ApiKey' or env 'OPENAI_API_KEY').")
    # Import lazily to avoid import-time failures
    from openai import OpenAI
    return OpenAI(api_key=key)

def extract_text(uploaded_file) -> str:
    """
    Extracts clean text from a FastAPI UploadFile-like object (PDF, DOCX, or TXT).
    Heavy libs are imported inside to avoid import-time side effects.
    """
    ext = os.path.splitext(uploaded_file.filename)[1].lower()
    file_bytes = uploaded_file.file.read()

    if ext == ".txt":
        return file_bytes.decode("utf-8", errors="ignore")

    if ext == ".docx":
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        out = []
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                out.append(t)
        for table in doc.tables:
            for row in table.rows:
                row_text = [c.text.strip() for c in row.cells if c.text.strip()]
                if row_text:
                    out.append(" | ".join(row_text))
        return "\n".join(out)

    if ext == ".pdf":
        import pdfplumber
        text = ""
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
        return text

    return ""

def get_tags(text: str, mode_prompt: str = "", num_tags: int = 8, mode: str = "Keywords") -> str:
    """
    Calls OpenAI to extract tags from text. Returns raw model string (parse with parse_tags()).
    """
    prompt = {
        "Keywords": f"Extract exactly {num_tags} concise keywords that summarize this document.",
        "Topics": f"List {num_tags} major themes or subjects covered in the document.",
        "Custom Prompt": mode_prompt or "Extract relevant tags.",
    }.get(mode, "Extract relevant tags.")

    client = _openai_client()
    resp = client.chat.completions.create(
        # Choose a lightweight model you actually have access to
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You extract short, relevant tags from documents."},
            {"role": "user", "content": f"{prompt}\n\n{text}"},
        ],
        temperature=0.3,
    )
    return resp.choices[0].message.content.strip()

def parse_tags(raw_text: str) -> List[str]:
    """
    Parses raw OpenAI output into a list of deduplicated tags.
    """
    lines = raw_text.splitlines()
    tags: List[str] = []
    for line in lines:
        line = re.sub(r"^[\s\-\•\*\d\.\)\(]+", "", line)
        parts = [t.strip() for t in line.split(",") if t.strip()]
        tags.extend(parts)

    seen = set()
    out: List[str] = []
    for t in tags:
        low = t.lower()
        if low not in seen:
            seen.add(low)
            out.append(t)
    return out
